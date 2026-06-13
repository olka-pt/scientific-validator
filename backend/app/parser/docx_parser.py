from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from backend.app.models.document import DocumentModel
from backend.app.models.paragraph import ParagraphModel

from backend.app.models.table import (
    TableModel,
    TableRowModel,
    TableCellModel,
)

class DocxParser:

    @staticmethod
    def parse(
        file_path: str
    ) -> DocumentModel:

        document = Document(file_path)

        page_count = None

        try:
            page_count = int(
                document.core_properties.pages
            )
        except Exception:
            page_count = None        

        has_page_numbering = DocxParser._has_page_numbering(
            document
        )

        paragraphs = []

        pending_page_break = False

        for index, paragraph in enumerate(document.paragraphs):

            text = paragraph.text.strip()

            current_has_page_break = DocxParser._has_page_break(
                paragraph
            )

            starts_new_page = (
                pending_page_break
                or DocxParser._has_page_break_before(paragraph)
            )

            if current_has_page_break:
                pending_page_break = True            

            if not text:
                continue

            pending_page_break = False

            first_run = (
                paragraph.runs[0]
                if paragraph.runs
                else None
            )

            font_family = None
            font_size = None
            bold = False
            italic = False

            if first_run:

                font_family = (
                    first_run.font.name
                )

                if first_run.font.size:
                    font_size = (
                        first_run.font.size.pt
                    )

                bold = bool(
                    first_run.bold
                )

                italic = bool(
                    first_run.italic
                )

            alignment = (
                DocxParser._parse_alignment(
                    paragraph.alignment
                )
            )

            paragraph_format = (
                paragraph.paragraph_format
            )

            line_spacing = (
                paragraph_format.line_spacing
            )

            if isinstance(line_spacing, float):

                line_spacing_value = (
                    round(line_spacing, 2)
                )

            elif line_spacing is not None:

                try:

                    line_spacing_value = (
                        round(
                            line_spacing.pt,
                            2
                        )
                    )

                except AttributeError:

                    line_spacing_value = None

            else:

                line_spacing_value = None

            first_line_indent = (
                DocxParser._convert_length(
                    paragraph_format.first_line_indent
                )
            )

            left_indent = (
                DocxParser._convert_length(
                    paragraph_format.left_indent
                )
            )

            right_indent = (
                DocxParser._convert_length(
                    paragraph_format.right_indent
                )
            )

            space_before = (
                DocxParser._convert_length(
                    paragraph_format.space_before
                )
            )

            space_after = (
                DocxParser._convert_length(
                    paragraph_format.space_after
                )
            )

            paragraph_model = (
                ParagraphModel(
                    text=text,
                    font_family=font_family,
                    font_size=font_size,
                    bold=bold,
                    italic=italic,
                    alignment=alignment,
                    starts_new_page=starts_new_page,                    

                    line_spacing=line_spacing_value,

                    first_line_indent=(
                        first_line_indent
                    ),

                    left_indent=left_indent,
                    right_indent=right_indent,

                    space_before=space_before,
                    space_after=space_after,

                    paragraph_index=index
                )
            )

            paragraphs.append(
                paragraph_model
            )

        section = document.sections[0]

        page_width = round(
            section.page_width.mm,
            2
        )

        page_height = round(
            section.page_height.mm,
            2
        )

        top_margin = round(
            section.top_margin.mm,
            2
        )

        bottom_margin = round(
            section.bottom_margin.mm,
            2
        )

        left_margin = round(
            section.left_margin.mm,
            2
        )

        right_margin = round(
            section.right_margin.mm,
            2
        )

        tables = []

        for table_index, table in enumerate(
            document.tables
        ):

            parsed_rows = []

            for row in table.rows:

                parsed_cells = []

                for cell in row.cells:

                    text = cell.text.strip()

                    font_size = None

                    for paragraph in cell.paragraphs:

                        for run in paragraph.runs:

                            if run.font.size:

                                font_size = (
                                    round(
                                        run.font.size.pt,
                                        1
                                    )
                                )

                                break

                        if font_size:
                            break

                    parsed_cells.append(
                        TableCellModel(
                            text=text,
                            font_size=font_size
                        )
                    )

                parsed_rows.append(
                    TableRowModel(
                        cells=parsed_cells
                    )
                )

            tables.append(
                TableModel(
                    rows=parsed_rows,
                    table_index=table_index
                )
            )      


        return DocumentModel(
            paragraphs=paragraphs,
            tables=tables,
            page_count=page_count,
            has_page_numbering=has_page_numbering,
            
            margins={
                "top": top_margin,
                "bottom": bottom_margin,

                "left": left_margin,
                "right": right_margin
            },

            page_size={
                "width": page_width,
                "height": page_height
            }
        )

    @staticmethod
    def _has_page_numbering(
        document
    ) -> bool:

        for section in document.sections:

            footer_xml = section.footer._element.xml.lower()

            if (
                "page" in footer_xml
                or "numpages" in footer_xml
            ):
                return True

            for paragraph in section.footer.paragraphs:

                text = paragraph.text.strip()

                if text.isdigit():
                    return True

        return False        

    @staticmethod
    def _has_page_break(
        paragraph
    ) -> bool:
        xml = paragraph._element.xml
        return 'w:type="page"' in xml

    @staticmethod
    def _has_page_break_before(
        paragraph
    ) -> bool:
        xml = paragraph._element.xml
        return "pageBreakBefore" in xml

    @staticmethod
    def _parse_alignment(
        alignment
    ) -> str | None:

        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify"
        }

        return alignment_map.get(
            alignment
        )

    @staticmethod
    def _convert_length(
        value
    ) -> float | None:

        if value is None:
            return None

        try:
            return round(
                value.cm,
                2
            )

        except Exception:
            return None